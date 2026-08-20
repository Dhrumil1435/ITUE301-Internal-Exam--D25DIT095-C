import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def generate_docx(docx_filename):
    doc = docx.Document()
    
    # Page Margins (1 inch = 0.75 for compact presentation)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("CHAROTAR UNIVERSITY OF SCIENCE AND TECHNOLOGY\nFaculty of Technology and Engineering — CSPIT-IT")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("ITUE301 — Advanced Web Development Frameworks\nOPEN-BOOK PRACTICAL EXAMINATION — SET A\nHospital Appointment System Report")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(15, 118, 110)
    
    # Metadata Table
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell_data = [
        [("Student Roll No: ", "24CSE001-A"), ("Tech Stack: ", "React + Express + MongoDB")],
        [("System Name: ", "MedCare Plus Hospital System"), ("Date: ", "August 20, 2026")]
    ]
    
    for r_idx, row in enumerate(cell_data):
        for c_idx, (label, val) in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_background(cell, "F1F5F9")
            p = cell.paragraphs[0]
            run_lbl = p.add_run(label)
            run_lbl.bold = True
            run_lbl.font.size = Pt(10)
            run_val = p.add_run(val)
            run_val.font.size = Pt(10)
            
    doc.add_paragraph() # Spacer
    
    # ----------------------------------------------------
    # SCREENSHOT 1 SECTION
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("Screenshot 1 — React Application")
    run_h1.font.color.rgb = RGBColor(15, 118, 110)
    
    p_desc1 = doc.add_paragraph("Demonstrates the Hospital Appointment System frontend running in the browser (React + Vite on http://localhost:5173). Shows navigation links, dashboard metrics, and status-styled appointment cards (confirmed, pending, cancelled).")
    p_desc1.runs[0].font.size = Pt(10.5)
    
    img_path1 = os.path.abspath("report_screenshots/screenshot1_react.png")
    if os.path.exists(img_path1):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture(img_path1, width=Inches(6.5))
        
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # SCREENSHOT 2 SECTION
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("Screenshot 2 — REST API Execution")
    run_h2.font.color.rgb = RGBColor(15, 118, 110)
    
    p_desc2 = doc.add_paragraph("Demonstrates a successful REST API request to the Express backend (GET http://localhost:5000/api/v1/doctors) returning HTTP status 200 OK and structured JSON data output.")
    p_desc2.runs[0].font.size = Pt(10.5)
    
    img_path2 = os.path.abspath("report_screenshots/screenshot2_rest_api.png")
    if os.path.exists(img_path2):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture(img_path2, width=Inches(6.5))
        
    doc.add_paragraph()
    
    # ----------------------------------------------------
    # SCREENSHOT 3 SECTION
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("Screenshot 3 — MongoDB Database & Documents")
    run_h3.font.color.rgb = RGBColor(15, 118, 110)
    
    p_desc3 = doc.add_paragraph("Demonstrates MongoDB Compass displaying the hospital_system database, appointments collection, and document schema fields (patientId, doctorId, status, timeSlot, date, reason).")
    p_desc3.runs[0].font.size = Pt(10.5)
    
    img_path3 = os.path.abspath("report_screenshots/screenshot3_mongodb.png")
    if os.path.exists(img_path3):
        p_img3 = doc.add_paragraph()
        p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img3.add_run().add_picture(img_path3, width=Inches(6.5))
        
    doc.save(docx_filename)
    print(f"Generated Word Document Report successfully: {docx_filename}")

if __name__ == "__main__":
    generate_docx("24CSE001_SetA_Report.docx")
    generate_docx("SetA_Report.docx")
